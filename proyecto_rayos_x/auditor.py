import time
from google import genai
from google.genai import types
from docx import Document
from PIL import Image

class AuditorSeguridad:
    def __init__(self, api_key: str):
        if not api_key:
            raise ValueError("API Key vacía.")
        self.client = genai.Client(api_key=api_key)

    def generar_dictamen_stream(self, pregunta: str, imagen: Image.Image, objeto: str, clasificacion: str):
        instrucciones_sistema = (
            "Eres un Auditor Senior de Seguridad Aeroportuaria y experto OACI. "
            "Responde de forma profesional, seria y pericial. "
            "IMPORTANTE: No uses negritas en ninguna parte de tu respuesta. "
            "Incluye un icono de escudo (🛡️) al inicio de tu informe técnico. "
            "Si detectas un arma, advierte al operador con firmeza. Ve directo al grano."
        )
        
        contenidos = [pregunta]
        if imagen is not None: contenidos.append(imagen)
            
        return self.client.models.generate_content_stream(
            model='gemini-2.5-flash',
            contents=contenidos,
            config=types.GenerateContentConfig(system_instruction=instrucciones_sistema, temperature=0.2)
        )

    def compilar_docx(self, analisis_modelo: str, objeto: str, clasificacion: str, umbral: float):
        doc = Document()
        
        doc.add_heading('SISTEMA NACIONAL DE CONTROL AEROPORTUARIO', 0)
        doc.add_paragraph('Acta de Inspección Técnica - Informe de Incidencia AVSEC')
        
        doc.add_heading('Datos de Identificación', level=1)
        doc.add_paragraph(f'ID Transacción: SR-RX-{int(time.time())}')
        doc.add_paragraph(f'Módulo de Visión: {objeto}')
        doc.add_paragraph('Nombre del Operador: ___________________________')
        doc.add_paragraph('ID de Empleado: ________________________________')
        
        doc.add_heading('Dictamen Técnico', level=1)
        texto_limpio = analisis_modelo.replace('**', '')
        doc.add_paragraph(texto_limpio)
        
        doc.add_heading('Marco Legal OACI', level=1)
        doc.add_paragraph('Este documento cumple con el Anexo 17 de la OACI sobre seguridad de la aviación civil.')
        
        doc.add_paragraph('\n\n')
        tabla_firmas = doc.add_table(rows=1, cols=2)
        tabla_firmas.cell(0, 0).text = '_____________________\nFirma Operador'
        tabla_firmas.cell(0, 1).text = '_____________________\nFirma Supervisor'
        
        nombre_archivo = "informe_incidencia_oaci.docx"
        doc.save(nombre_archivo)
        return nombre_archivo